# This is a LLM chatbot that allows architects to develop a scope of work for AEC contracts.
# Each interaction with the chatbot refines the scope of work, ultimately producing a comprehensive
# list of deliverables,
# along with associated scope items for each deliverable.

import pandas as pd
import streamlit as st
from llm_calls import classify_and_get_context, update_categories_list
from scope_visualizer import display_scope_of_work
from ui_styles import apply_custom_styles
from docx import Document
import io

HEIGHT = 500

default_scope_of_work_dict_file = "default_scope_of_work.txt"
default_scope_of_work_string = open(default_scope_of_work_dict_file, "r").read()
default_scope_of_work = eval(default_scope_of_work_string)

default_categories_list_file = "default_categories_list.txt"
default_categories_list_string = open(default_categories_list_file, "r").read()
st.session_state["FULL_CATEGORIES_LIST"] = eval(default_categories_list_string)

default_assumptions_and_exclusions_file = "default_assumptions_and_exclusions.txt"
default_assumptions_and_exclusions_string = open(default_assumptions_and_exclusions_file, "r").read()
default_assumptions_and_exclusions = eval(default_assumptions_and_exclusions_string)
st.session_state.setdefault("ASSUMPTIONS_AND_EXCLUSIONS", default_assumptions_and_exclusions)

print("\n" * 5)
print("Starting AEC Contract Assistant...")

st.set_page_config(
    page_title="ContractCadence",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Apply custom CSS styling
# apply_custom_styles()

# Build the chat interface
st.title("ContractCadence 🤖")
doc_column, chat_column = st.columns([3, 3])

# The chat window allows the user to have a conversation with the AI assistant
# This conversation would generate a dictionary of deliverables, and associated lists of scope items

st.session_state.setdefault("messages", [])
st.session_state.setdefault("scope_of_work", default_scope_of_work)
st.session_state.setdefault("conversation_history", [])

st.session_state.setdefault("first_run", True)

print("Streamlit session state initialized.")
print(f"Scope of work: {st.session_state.scope_of_work}")
print(f"Assumptions and Exclusions: {st.session_state['ASSUMPTIONS_AND_EXCLUSIONS']}")

with st.sidebar:
    st.markdown("##### Upload Reference Documents 📄")
    uploaded_files = st.file_uploader(
        "Upload architectural plans, contracts, or other reference documents to assist the AI in understanding your project requirements. Note: This feature is currently disabled.",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt"],
    )
    if uploaded_files:
        st.markdown("**Uploaded Files:**")
        for uploaded_file in uploaded_files:
            st.markdown(f"- {uploaded_file.name}")
            
# The scope window displays the current scope of work being developed
# It shows a list of deliverables, each with associated scope items
with doc_column:
    scope_tab, exclusions_tab = st.tabs(["Scope of Work 📋", "Assumptions and Exclusions ❌"])

    with scope_tab:
        tabs = st.tabs(["Markdown View", "Table View (Manual Edit)"])
        
        with tabs[1]:
            # Use DEFAULT_SCOPE_OF_WORK for testing
            scope_to_display = st.session_state.scope_of_work
            display_scope_of_work(scope_to_display, height=HEIGHT)

        with tabs[0]:
            with st.container(height=HEIGHT, border=True):
                # Display the scope of work in markdown format
                scope_of_work = st.session_state.scope_of_work

                markdown_lines = []
                try:
                    for phase, disciplines in scope_of_work.items():
                        markdown_lines.append(f"##### {phase}")
                        try:
                            for discipline, items in disciplines.items():
                                markdown_lines.append(f"###### {discipline}")
                                for item in items:
                                    markdown_lines.append(f"- {item}")
                        except Exception as e:
                            continue
                except Exception as e:
                    markdown_lines.append("No scope items yet.")
                markdown_lines.append("\n")
                markdown_lines.append("\* This scope of work was generated with the assistance of ContractCadence, an AI-powered AEC contract assistant.")

                markdown_text = "\n".join(markdown_lines)
                st.markdown(markdown_text)

        cols = st.columns(3)
        with cols[0]:
            # Add a download button for the scope of work
            scope_of_work_df = pd.DataFrame(st.session_state.scope_of_work)
            download_button = st.download_button(
                label="Download scope of work as CSV",
                data=scope_of_work_df.to_csv(index=False),
                file_name="scope_of_work.csv",
                mime="text/csv",
                width="stretch",
            )
        with cols[1]:
            # Add a button to download the markdown view of the scope of work as a .md file
            doc = Document()
            scope_of_work = st.session_state.scope_of_work
            try:
                for phase, disciplines in scope_of_work.items():
                    doc.add_heading(phase, level=1)
                    try:
                        for discipline, items in disciplines.items():
                            doc.add_heading(discipline, level=2)
                            for item in items:
                                doc.add_paragraph(item, style='List Bullet')
                    except Exception as e:
                        continue
            except Exception as e:
                pass
            docx_content = io.BytesIO()
            doc.save(docx_content)
            docx_content.seek(0)

            download_button = st.download_button(
                label="Download scope of work as DOCX",
                data=docx_content,
                file_name="scope_of_work.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width="stretch",
            )
        with cols[2]:
            download_button = st.button(
                label="Download categories list as CSV",
                width="stretch",
            )

            if download_button:
                with st.spinner("Updating categories list, this may take a moment..."):
                    update_categories_list()

    with exclusions_tab:        
        tabs = st.tabs(["Markdown View", "Table View (Manual Edit)"])
        with tabs[1]:
            # Display the assumptions and exclusions in a table view for manual editing
            assumptions_and_exclusions = st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"]
            
            disciplines = assumptions_and_exclusions.keys()
            discipline_col_rows = []
            assumptions_and_exclusions_col_rows = []
            for discipline in disciplines:
                items = assumptions_and_exclusions[discipline]
                for item in items:
                    discipline_col_rows.append(discipline)
                    assumptions_and_exclusions_col_rows.append(item)

            df_ae = pd.DataFrame({
                "Discipline": discipline_col_rows,
                "Assumptions & Exclusions": assumptions_and_exclusions_col_rows
            })

            edited_df_ae = st.data_editor(df_ae, height=HEIGHT, width="stretch", hide_index=True, num_rows="dynamic")
            edited_dict_ae = edited_df_ae.to_dict(orient='records')
            # Convert back to nested dictionary
            st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"] = {}
            for row in edited_dict_ae:
                discipline = row["Discipline"]
                item = row["Assumptions & Exclusions"]
                if discipline not in st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"]:
                    st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"][discipline] = []
                st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"][discipline].append(item)

        with tabs[0]:
            with st.container(height=HEIGHT, border=True):
                # Display the assumptions and exclusions in a markdown view
                markdown_lines_ae = []
                for discipline, items in st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"].items():
                    markdown_lines_ae.append(f"##### {discipline}")
                    for item in items:
                        markdown_lines_ae.append(f"- {item}")
                markdown_text_ae = "\n".join(markdown_lines_ae)
                st.markdown(markdown_text_ae)
        
        cols = st.columns(2)
        with cols[0]:
            # Add a download button for the assumptions and exclusions as CSV
            ae_discipline_col_rows = []
            ae_item_col_rows = []
            for discipline, items in st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"].items():
                for item in items:
                    ae_discipline_col_rows.append(discipline)
                    ae_item_col_rows.append(item)

            ae_df = pd.DataFrame({
                "Discipline": ae_discipline_col_rows,
                "Assumptions & Exclusions": ae_item_col_rows
            })

            download_button_ae = st.download_button(
                label="Download assumptions and exclusions as CSV",
                data=ae_df.to_csv(index=False),
                file_name="assumptions_and_exclusions.csv",
                mime="text/csv",
                width="stretch",
            )
        with cols[1]:
            # Add a button to download the assumptions and exclusions as a DOCX file
            doc_ae = Document()
            doc_ae.add_heading("Assumptions and Exclusions", level=0)
            for discipline, items in st.session_state["ASSUMPTIONS_AND_EXCLUSIONS"].items():
                doc_ae.add_heading(discipline, level=1)
                for item in items:
                    doc_ae.add_paragraph(item, style='List Bullet')
            docx_content_ae = io.BytesIO()
            doc_ae.save(docx_content_ae)
            docx_content_ae.seek(0)

            download_button_ae = st.download_button(
                label="Download assumptions and exclusions as DOCX",
                data=docx_content_ae,
                file_name="assumptions_and_exclusions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width="stretch",
            )


with chat_column:
    st.markdown("##### Chat with your AEC Contract Assistant 💬")
    message_container = st.container(height=HEIGHT + 90, border=True)

    if st.session_state.messages == []:
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": "Hello! I'm your AEC Contract Assistant. Describe your project and requirements, and I'll help you build a comprehensive scope of work."
            }
        )

    with message_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    if prompt := st.chat_input("Describe your project and requirements..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with message_container:
            with st.chat_message("user"):
                st.markdown(prompt)

            # Call LLM with conversation history for context, with streaming enabled
            response_generator = classify_and_get_context(prompt)

            # Display streaming response
            with st.chat_message("assistant"):
                response_placeholder = st.empty()
                full_response = ""
                
                # Stream the response
                for chunk in response_generator:
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")
                
                # Final update without cursor
                response_placeholder.markdown(full_response)

        # Update conversation history with user message and assistant response
        st.session_state.conversation_history.append({"role": "user", "content": prompt})
        st.session_state.conversation_history.append({"role": "assistant", "content": full_response})

        st.session_state.messages.append({"role": "assistant", "content": full_response})

        print("Rerunning Streamlit app.")
        st.rerun()
            
str_list = []
str_list.append("This information should not be considered legal advice. Consult a qualified attorney for legal matters.\n")
str_list.append("Github Repository: https://github.com/sclebow/aectech25-nyc-legal-chat/ \n")
str_list.append("Contributers: ")
contributers_list = ["Scott Lebow, https://www.linkedin.com/in/sclebow/",
                     "Chu Ding, https://www.linkedin.com/in/chuding/",
                     "Yufei Wang, https://www.linkedin.com/in/yufei-wang-faye/",
                     "Douglas Kim, https://www.linkedin.com/in/dkim19/",
                     "Janez Mikec, https://www.linkedin.com/in/janezmikec/"]

# Randomize contributers order
import random
random.shuffle(contributers_list)

for contributer in contributers_list:
    str_list.append(f"- {contributer}")

disclaimer_text = "\n".join(str_list)
with st.expander("Disclaimer & Contributers"):
    st.markdown(disclaimer_text)