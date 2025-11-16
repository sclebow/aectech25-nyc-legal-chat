import requests
import viktor as vkt

# AEC Data Model GraphQL endpoint
AEC_GRAPHQL_URL = "https://developer.api.autodesk.com/aec/graphql"


def execute_graphql(
    query: str, token: str, region: str, variables: dict = None, timeout: int = 30
):
    """
    Execute a GraphQL query against the Autodesk AEC Data Model API.

    Args:
        query: GraphQL query string
        token: OAuth2 access token
        region: Region identifier (e.g., 'US', 'EMEA')
        variables: Optional dictionary of GraphQL variables
        timeout: Request timeout in seconds

    Returns:
        Dictionary containing the response data
    """
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Region": region,
    }
    payload = {"query": query, "variables": variables or {}}
    resp = requests.post(
        AEC_GRAPHQL_URL, headers=headers, json=payload, timeout=timeout
    )

    if resp.status_code != 200:
        raise RuntimeError(f"HTTP {resp.status_code}: {resp.text}")

    body = resp.json()
    if body.get("errors"):
        raise RuntimeError(f"GraphQL errors: {body['errors']}")

    return body.get("data", {})


class Parametrization(vkt.Parametrization):
    """Application input parameters organized in steps."""

    # Step 1: Autodesk Files
    step_1 = vkt.Step(
        "Autodesk Files",
        views=["view_autodesk_model"],
        description="Select the Autodesk files to analyze",
        next_label="Next: Define Contract Scope",
    )

    step_1.autodesk_file = vkt.AutodeskFileField(
        "Select Autodesk Structural File",
        oauth2_integration="autodesk-integration",
        description="Select a structural file from Autodesk Construction Cloud to view and analyze",
        flex=60,
    )

    step_1.autodesk_file_electrical = vkt.AutodeskFileField(
        "Select Autodesk Electrical File",
        oauth2_integration="autodesk-integration",
        description="Select an electrical file from Autodesk Construction Cloud to view and analyze",
        flex=60,
    )

    # Step 2: Contract Scope
    step_2 = vkt.Step(
        "Contract Scope",
        views=[
            "view_category_summary",
            "view_colored_categories",
            "view_category_data",
        ],
        description="Define the categories that should be present in the model",
        previous_label="Back: Autodesk Files",
        next_label="Next: Download Report",
    )

    step_2.csv_file = vkt.FileField(
        "Contract Scope",
        file_types=[".csv"],
        description="Upload a CSV file with categories in rows (no header)",
        flex=60,
    )

    step_2.load_from_csv = vkt.SetParamsButton(
        "Load Contract Scope",
        method="load_categories_from_csv",
        description="Populate the Contract Scope with categories from the uploaded CSV file",
        flex=60,
    )

    step_2.required_categories = vkt.DynamicArray(
        "Contract Scope",
        description="Add categories that should be present in the model with custom colors",
    )
    step_2.required_categories.category = vkt.OptionField(
        "Category",
        options=[
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ],
    )
    step_2.required_categories.color = vkt.ColorField(
        "Highlight Color", default=vkt.Color(0, 255, 0)
    )

    # Step 3: Download Report
    step_3 = vkt.Step(
        "Download Report",
        views=["view_category_summary"],
        description="Download a comprehensive report with the category summary",
        previous_label="Back: Contract Scope",
    )

    step_3.download_report = vkt.DownloadButton(
        "Download Contract Compliance Report",
        method="download_contract_compliance_report",
        description="Download a Word document showing how the model complies with the contract scope",
        flex=60,
    )


class Controller(vkt.Controller):
    """Main application controller."""

    parametrization = Parametrization

    def load_categories_from_csv(self, params, **kwargs):
        """
        Load categories from the uploaded CSV file and populate the Dynamic Array.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            SetParamsResult with updated required_categories
        """
        import csv
        import random

        if not params.step_2.csv_file:
            raise vkt.UserError("Please upload a CSV file first")

        # Read the CSV file
        csv_file = params.step_2.csv_file.file
        categories = []

        try:
            with csv_file.open() as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    # Skip empty rows
                    if row and row[0].strip():
                        category_name = row[0].strip()
                        categories.append(category_name)
        except Exception as e:
            raise vkt.UserError(f"Failed to read CSV file: {str(e)}")

        if not categories:
            raise vkt.UserError("No categories found in the CSV file")

        # Generate random colors for each category
        new_categories = []
        for category in categories:
            # Generate a random color
            r = random.randint(50, 255)
            g = random.randint(50, 255)
            b = random.randint(50, 255)

            new_categories.append({"category": category, "color": vkt.Color(r, g, b)})

        # Return SetParamsResult to update the required_categories field
        return vkt.SetParamsResult({"step_2": {"required_categories": new_categories}})

    @vkt.AutodeskView("3D Model Viewer", duration_guess=5)
    def view_autodesk_model(self, params, **kwargs):
        """
        Display the selected Autodesk file in the 3D viewer.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            AutodeskResult containing the model to display
        """
        if not params.step_1.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Return the Autodesk viewer result
        return vkt.AutodeskResult(params.step_1.autodesk_file, access_token=token)

    @vkt.TableView("Category Summary", duration_guess=10)
    def view_category_summary(self, params, **kwargs):
        """
        Display a summary table that shows the same categories as the dropdown list
        and cross-checks whether they are present in the models (structural and/or electrical).

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            TableResult showing which categories from the dropdown are present in the models
        """
        if (
            not params.step_1.autodesk_file
            and not params.step_1.autodesk_file_electrical
        ):
            raise vkt.UserError(
                "Please select at least one Autodesk file (structural or electrical)"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Collect category counts from both files
        structural_counts = {}
        electrical_counts = {}

        # Extract required categories from dynamic array
        required_categories = set(
            row["category"] for row in params.step_2.required_categories
        )

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        # Query to get all distinct categories in a model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        # Fetch from structural file if provided
        if params.step_1.autodesk_file:
            vkt.progress_message(
                "Fetching categories from structural file...", percentage=20
            )
            try:
                region = params.step_1.autodesk_file.get_region(token)
                group_id = (
                    params.step_1.autodesk_file.get_aec_data_model_element_group_id(
                        token
                    )
                )

                variables = {"elementGroupId": group_id, "limit": 1000}
                data = execute_graphql(query, token, region, variables)
                block = data.get("distinctPropertyValuesInElementGroupByName") or {}
                results_list = block.get("results") or []

                for r in results_list:
                    values = r.get("values") or []
                    for v in values:
                        category_name = v.get("value", "")
                        element_count = v.get("count", 0)
                        if category_name:
                            structural_counts[category_name] = element_count

            except Exception as e:
                vkt.UserMessage.warning(
                    f"Failed to fetch categories from structural file: {str(e)}"
                )

        # Fetch from electrical file if provided
        if params.step_1.autodesk_file_electrical:
            vkt.progress_message(
                "Fetching categories from electrical file...", percentage=50
            )
            try:
                region = params.step_1.autodesk_file_electrical.get_region(token)
                group_id = params.step_1.autodesk_file_electrical.get_aec_data_model_element_group_id(
                    token
                )

                variables = {"elementGroupId": group_id, "limit": 1000}
                data = execute_graphql(query, token, region, variables)
                block = data.get("distinctPropertyValuesInElementGroupByName") or {}
                results_list = block.get("results") or []

                for r in results_list:
                    values = r.get("values") or []
                    for v in values:
                        category_name = v.get("value", "")
                        element_count = v.get("count", 0)
                        if category_name:
                            electrical_counts[category_name] = element_count

            except Exception as e:
                vkt.UserMessage.warning(
                    f"Failed to fetch categories from electrical file: {str(e)}"
                )

        vkt.progress_message("Preparing category summary...", percentage=80)

        # Prepare table data with visual indicators
        table_data = []
        for category_name in all_categories:
            # Get element counts from both files
            structural_count = structural_counts.get(category_name, 0)
            electrical_count = electrical_counts.get(category_name, 0)
            total_count = structural_count + electrical_count

            # Check if category is in any model
            in_model = total_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Build element count display with breakdown
            if params.step_1.autodesk_file and params.step_1.autodesk_file_electrical:
                count_display = (
                    f"{total_count} (S:{structural_count}, E:{electrical_count})"
                )
            elif params.step_1.autodesk_file:
                count_display = f"{structural_count}"
            else:
                count_display = f"{electrical_count}"

            # Determine status symbol and description
            if in_contract and in_model:
                status_symbol = "✓"
                status_text = "Present in contract and model(s)"
                status_color = vkt.Color(0, 128, 0)  # Green
            elif in_contract and not in_model:
                status_symbol = "✗"
                status_text = "In contract but not in model(s)"
                status_color = vkt.Color(255, 165, 0)  # Orange
            elif not in_contract and in_model:
                status_symbol = "✗"
                status_text = "Missing in the contract"
                status_color = vkt.Color(255, 0, 0)  # Red
            else:  # not in_contract and not in_model
                status_symbol = "✗"
                status_text = "Not in contract, not in model(s)"
                status_color = vkt.Color(128, 128, 128)  # Gray

            # Create colored cells for better visualization
            status_cell = vkt.TableCell(
                status_symbol, text_color=status_color, text_style="bold"
            )

            table_data.append([category_name, status_cell, count_display, status_text])

        # Define column headers
        column_headers = [
            vkt.TableHeader("Category", align="left"),
            vkt.TableHeader("Status", align="center"),
            vkt.TableHeader("Element Count", align="right"),
            vkt.TableHeader("Description", align="left"),
        ]

        return vkt.TableResult(
            table_data, column_headers=column_headers, enable_sorting_and_filtering=True
        )

    @vkt.WebView("Colored Category View", duration_guess=15)
    def view_colored_categories(self, params, **kwargs):
        """
        Display the Autodesk model with categories highlighted in custom colors
        based on the dynamic field selections.
        """
        if not params.step_1.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get the URN from the Autodesk file and encode it properly
        autodesk_file = params.step_1.autodesk_file
        region = autodesk_file.get_region(token)
        group_id = autodesk_file.get_aec_data_model_element_group_id(token)

        # Get the latest version URN and encode it like in your working example
        latest_version = autodesk_file.get_latest_version(token)
        urn = latest_version.urn
        import base64

        urn_bs64 = base64.urlsafe_b64encode(urn.encode()).decode().rstrip("=")

        vkt.progress_message(
            "Fetching element external IDs for selected categories...", percentage=20
        )

        # Build a list of external IDs with their colors for each category
        external_ids_with_colors = []

        for row in params.step_2.required_categories:
            category_name = row["category"]
            color = row["color"]

            # Convert VIKTOR Color to hex format
            color_hex = color.hex

            vkt.progress_message(f"Fetching {category_name} elements...", percentage=30)

            # GraphQL query to get element IDs and their external IDs
            query = """
            query CategoryElements($elementGroupId: ID!, $rsqlFilter: String!, $pagination: PaginationInput) {
            elementsByElementGroup(
                elementGroupId: $elementGroupId,
                filter: { query: $rsqlFilter },
                pagination: $pagination
            ) {
                pagination { cursor pageSize }
                results {
                id
                name
                alternativeIdentifiers {
                    externalElementId
                }
                }
            }
            }
            """

            # Construct RSQL filter for this category
            rsql_filter = f"property.name.category=='{category_name}' and 'property.name.Element Context'==Instance"

            # Fetch all elements with pagination
            cursor = None
            limit = 100

            while True:
                variables = {
                    "elementGroupId": group_id,
                    "rsqlFilter": rsql_filter,
                    "pagination": {"limit": limit}
                    if not cursor
                    else {"cursor": cursor, "limit": limit},
                }

                try:
                    data = execute_graphql(query, token, region, variables)
                    block = data.get("elementsByElementGroup", {}) or {}
                    page_results = block.get("results", []) or []

                    # Collect external IDs with their colors
                    for element in page_results:
                        # Get External ID from alternativeIdentifiers
                        alt_ids = element.get("alternativeIdentifiers", {})
                        external_id = alt_ids.get("externalElementId")

                        # If External ID found, add it with its color
                        if external_id:
                            # Create a single-key object as expected by the viewer script
                            external_ids_with_colors.append({external_id: color_hex})

                    # Check pagination
                    page = block.get("pagination", {}) or {}
                    new_cursor = page.get("cursor")

                    # Stop on empty cursor, repeated cursor, or empty page
                    if not new_cursor or new_cursor == cursor or len(page_results) == 0:
                        break

                    cursor = new_cursor

                except Exception as e:
                    vkt.UserMessage.warning(
                        f"Could not fetch elements for category '{category_name}': {str(e)}"
                    )
                    break

        vkt.progress_message("Preparing viewer...", percentage=80)

        # Convert Python list to JSON string for JavaScript
        import json

        external_ids_json = json.dumps(external_ids_with_colors)

        # Use the same HTML template approach as your working example
        html_template = """<!DOCTYPE html>
    <html>
    <head>
    <meta charset="utf-8" />
    <title>APS Viewer - Colored Categories</title>
    <link rel="stylesheet" href="https://developer.api.autodesk.com/modelderivative/v2/viewers/7.*/style.min.css" type="text/css">
    <script src="https://developer.api.autodesk.com/modelderivative/v2/viewers/7.*/viewer3D.min.js"></script>
    <style>
        html, body, #apsViewerDiv { width: 100%; height: 100%; margin: 0; }
    </style>
    </head>
    <body>
    <div id="apsViewerDiv"></div>

    <script>
    // Injected external IDs from backend
    var EXTERNAL_IDS = EXTERNAL_IDS_PLACEHOLDER;
    // Replace these three in your backend
    var ACCESS_TOKEN = 'APS_TOKEN_PLACEHOLDER';
    var DOCUMENT_URN = 'urn:URN_PLACEHOLDER';
    
    console.log('External IDs:', EXTERNAL_IDS);

    // Disable analytics for sandbox iframes
    try { Autodesk.Viewing.Private.analytics.optOut(); } catch (e) {}

    let viewer = null;
    let modelLoaded = null;

    Autodesk.Viewing.Initializer(
        { env: 'AutodeskProduction2', api: 'streamingV2', accessToken: ACCESS_TOKEN },
        function () {
        const container = document.getElementById('apsViewerDiv');
        viewer = new Autodesk.Viewing.GuiViewer3D(container, { disableBimWalkInfoIcon: true });
        viewer.start();
        console.log('Viewer started');

        if (!DOCUMENT_URN) {
            console.error('Missing URN');
            return;
        }

        Autodesk.Viewing.Document.load(
            DOCUMENT_URN,
            function onSuccess(doc) {
            const node = doc.getRoot().getDefaultGeometry();
            if (!node) { console.warn('No default geometry'); return; }
            viewer.loadDocumentNode(doc, node, { keepCurrentModels: false }).then(function (model) {
                modelLoaded = model;
                console.log('Model loaded');

                // Apply filtering and coloring
                modelLoaded.getExternalIdMapping(
                function onMap(map) {
                    if (!map) {
                    console.warn('externalId map not available');
                    return;
                    }
                    
                    const dbIds = [];
                    const extColorMap = buildExternalColorMap(EXTERNAL_IDS);
                    const rev = {}; // dbId -> externalId
                    
                    for (const extId in map) {
                    if (Object.prototype.hasOwnProperty.call(map, extId)) {
                        rev[ map[extId] ] = extId;
                    }
                    }
                    
                    // Collect dbIds for all external IDs in EXTERNAL_IDS
                    for (const extId in extColorMap) {
                    if (map[extId]) {
                        dbIds.push(map[extId]);
                    }
                    }
                    
                    if (dbIds.length === 0) {
                    console.info('No dbIds match EXTERNAL_IDS');
                    return;
                    }
                    
                    viewer.clearThemingColors();
                    viewer.isolate(dbIds);
                    viewer.fitToView(dbIds);
                    
                    // Apply colors
                    const THREE_REF = (Autodesk && Autodesk.Viewing && Autodesk.Viewing.Private && Autodesk.Viewing.Private.THREE) || window.THREE;
                    const defaultV4 = colorStringToVec4('green', 0.85, THREE_REF);
                    
                    for (let i = 0; i < dbIds.length; i++) {
                    viewer.setThemingColor(dbIds[i], defaultV4, modelLoaded, false);
                    }
                    
                    // Apply overrides from extColorMap
                    for (let i = 0; i < dbIds.length; i++) {
                    const dbId = dbIds[i];
                    const ext = rev[dbId];
                    if (ext && extColorMap[ext]) {
                        const v4 = colorStringToVec4(extColorMap[ext], 0.95, THREE_REF);
                        viewer.setThemingColor(dbId, v4, modelLoaded, false);
                    }
                    }
                    viewer.impl.invalidate(true, true, true);
                },
                function onErr(err) {
                    console.error('getExternalIdMapping failed', err);
                }
                );
            });
            },
            function onFailure(code, message) {
            console.error('Document load failed:', code, message);
            }
        );
        }
    );

    // Helper functions
    function buildExternalColorMap(list) {
        const map = Object.create(null);
        if (!Array.isArray(list)) return map;
        for (let i = 0; i < list.length; i++) {
        const obj = list[i];
        if (obj && typeof obj === 'object') {
            const keys = Object.keys(obj);
            if (keys.length === 1) {
            const extId = String(keys[0]);
            const color = String(obj[extId] || '').trim();
            if (extId && color) map[extId] = color;
            }
        }
        }
        return map;
    }

    function colorStringToVec4(str, alphaDefault, THREE_REF) {
        let a = typeof alphaDefault === 'number' ? alphaDefault : 0.85;
        if (!str) return new THREE_REF.Vector4(0, 1, 0, a);

        let s = String(str).trim().toLowerCase();
        
        // #rrggbb
        if (s.startsWith('#') && s.length === 7) {
        const r = parseInt(s.slice(1, 3), 16) / 255;
        const g = parseInt(s.slice(3, 5), 16) / 255;
        const b = parseInt(s.slice(5, 7), 16) / 255;
        return new THREE_REF.Vector4(r, g, b, a);
        }

        // fallback to green
        return new THREE_REF.Vector4(0, 1, 0, a);
    }
    </script>
    </body>
    </html>"""

        # Use placeholder replacement like in your working example
        html = html_template.replace("APS_TOKEN_PLACEHOLDER", token)
        html = html.replace("URN_PLACEHOLDER", urn_bs64)
        html = html.replace("EXTERNAL_IDS_PLACEHOLDER", external_ids_json)

        return vkt.WebResult(html=html)

    @vkt.DataView("Category Data Summary", duration_guess=10)
    def view_category_data(self, params, **kwargs):
        """
        Display a data summary showing which categories from the dropdown are present in the model.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            DataResult with category status information
        """
        if not params.step_1.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get region and AEC Data Model element group ID from the Autodesk file
        region = params.step_1.autodesk_file.get_region(token)
        group_id = params.step_1.autodesk_file.get_aec_data_model_element_group_id(
            token
        )

        # Extract required categories from dynamic array
        required_categories = set(
            row["category"] for row in params.step_2.required_categories
        )

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        vkt.progress_message("Fetching category counts from model...", percentage=10)

        # Query to get all distinct categories in the model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        variables = {
            "elementGroupId": group_id,
            "limit": 1000,  # High limit to get all categories
        }

        try:
            data = execute_graphql(query, token, region, variables)
            block = data.get("distinctPropertyValuesInElementGroupByName") or {}
            results_list = block.get("results") or []

            # Create a dictionary of category counts from the model
            model_category_counts = {}
            for r in results_list:
                values = r.get("values") or []
                for v in values:
                    category_name = v.get("value", "")
                    element_count = v.get("count", 0)
                    if category_name:
                        model_category_counts[category_name] = element_count

        except Exception as e:
            raise vkt.UserError(f"Failed to fetch categories from model: {str(e)}")

        vkt.progress_message("Preparing category data summary...", percentage=80)

        # Create main data group
        main_group = vkt.DataGroup()

        # Add summary statistics
        total_categories = len(all_categories)
        categories_in_model = sum(
            1 for cat in all_categories if model_category_counts.get(cat, 0) > 0
        )
        categories_in_contract = len(required_categories)
        categories_matched = sum(
            1 for cat in required_categories if model_category_counts.get(cat, 0) > 0
        )

        summary_group = vkt.DataGroup(
            vkt.DataItem("Total Categories", total_categories),
            vkt.DataItem("Categories in Model", categories_in_model),
            vkt.DataItem("Categories in Contract", categories_in_contract),
            vkt.DataItem(
                "Contract Categories Found",
                categories_matched,
                status=vkt.DataStatus.SUCCESS
                if categories_matched == categories_in_contract
                else vkt.DataStatus.WARNING,
            ),
        )
        main_group.add(vkt.DataItem("Summary", subgroup=summary_group))

        # Add category details grouped by status
        present_group = vkt.DataGroup()
        missing_from_model_group = vkt.DataGroup()
        missing_from_contract_group = vkt.DataGroup()
        not_applicable_group = vkt.DataGroup()

        for category_name in all_categories:
            # Check if category is in the model
            element_count = model_category_counts.get(category_name, 0)
            in_model = element_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Categorize and add to appropriate group
            if in_contract and in_model:
                # Present in both contract and model - SUCCESS
                present_group.add(
                    vkt.DataItem(
                        category_name,
                        element_count,
                        suffix="elements",
                        status=vkt.DataStatus.SUCCESS,
                        status_message="✓ Present in contract and model",
                    )
                )
            elif in_contract and not in_model:
                # In contract but not in model - ERROR
                missing_from_model_group.add(
                    vkt.DataItem(
                        category_name,
                        "0",
                        suffix="elements",
                        status=vkt.DataStatus.ERROR,
                        status_message="✗ In contract but not in model",
                    )
                )
            elif not in_contract and in_model:
                # In model but missing from contract - WARNING
                missing_from_contract_group.add(
                    vkt.DataItem(
                        category_name,
                        element_count,
                        suffix="elements",
                        status=vkt.DataStatus.WARNING,
                        status_message="✗ Missing in the contract",
                    )
                )
            else:
                # Not in contract and not in model - INFO
                not_applicable_group.add(
                    vkt.DataItem(
                        category_name,
                        "0",
                        suffix="elements",
                        status=vkt.DataStatus.INFO,
                        status_message="Not in contract, not in model",
                    )
                )

        # Add grouped categories to main group
        if len(present_group) > 0:
            main_group.add(
                vkt.DataItem("✓ Present (Contract & Model)", subgroup=present_group)
            )

        if len(missing_from_model_group) > 0:
            main_group.add(
                vkt.DataItem("✗ Missing from Model", subgroup=missing_from_model_group)
            )

        if len(missing_from_contract_group) > 0:
            main_group.add(
                vkt.DataItem(
                    "⚠ Missing from Contract", subgroup=missing_from_contract_group
                )
            )

        if len(not_applicable_group) > 0:
            main_group.add(
                vkt.DataItem("○ Not Applicable", subgroup=not_applicable_group)
            )

        return vkt.DataResult(main_group)

    def download_contract_compliance_report(self, params, **kwargs):
        """
        Generate and download a Word document showing how the model complies with the contract scope.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            DownloadResult with the Word document
        """
        import io
        from datetime import datetime

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        if (
            not params.step_1.autodesk_file
            and not params.step_1.autodesk_file_electrical
        ):
            raise vkt.UserError(
                "Please select at least one Autodesk file (structural or electrical)"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Collect category counts from both files
        structural_counts = {}
        electrical_counts = {}

        # Extract required categories from dynamic array
        required_categories = set(
            row["category"] for row in params.step_2.required_categories
        )

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        # Query to get all distinct categories in a model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        # Fetch from structural file if provided
        if params.step_1.autodesk_file:
            vkt.progress_message(
                "Fetching categories from structural file...", percentage=10
            )
            try:
                region = params.step_1.autodesk_file.get_region(token)
                group_id = (
                    params.step_1.autodesk_file.get_aec_data_model_element_group_id(
                        token
                    )
                )

                variables = {"elementGroupId": group_id, "limit": 1000}
                data = execute_graphql(query, token, region, variables)
                block = data.get("distinctPropertyValuesInElementGroupByName") or {}
                results_list = block.get("results") or []

                for r in results_list:
                    values = r.get("values") or []
                    for v in values:
                        category_name = v.get("value", "")
                        element_count = v.get("count", 0)
                        if category_name:
                            structural_counts[category_name] = element_count

            except Exception as e:
                vkt.UserMessage.warning(
                    f"Failed to fetch categories from structural file: {str(e)}"
                )

        # Fetch from electrical file if provided
        if params.step_1.autodesk_file_electrical:
            vkt.progress_message(
                "Fetching categories from electrical file...", percentage=30
            )
            try:
                region = params.step_1.autodesk_file_electrical.get_region(token)
                group_id = params.step_1.autodesk_file_electrical.get_aec_data_model_element_group_id(
                    token
                )

                variables = {"elementGroupId": group_id, "limit": 1000}
                data = execute_graphql(query, token, region, variables)
                block = data.get("distinctPropertyValuesInElementGroupByName") or {}
                results_list = block.get("results") or []

                for r in results_list:
                    values = r.get("values") or []
                    for v in values:
                        category_name = v.get("value", "")
                        element_count = v.get("count", 0)
                        if category_name:
                            electrical_counts[category_name] = element_count

            except Exception as e:
                vkt.UserMessage.warning(
                    f"Failed to fetch categories from electrical file: {str(e)}"
                )

        vkt.progress_message("Generating Word document...", percentage=60)

        # Create Word document
        doc = Document()

        # Add title
        title = doc.add_heading("Contract Compliance Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Add file information
        if params.step_1.autodesk_file and params.step_1.autodesk_file_electrical:
            doc.add_paragraph(
                f"Structural File: {params.step_1.autodesk_file.url.split('/')[-1]}"
            )
            doc.add_paragraph(
                f"Electrical File: {params.step_1.autodesk_file_electrical.url.split('/')[-1]}"
            )
        elif params.step_1.autodesk_file:
            doc.add_paragraph(
                f"Structural File: {params.step_1.autodesk_file.url.split('/')[-1]}"
            )
        else:
            doc.add_paragraph(
                f"Electrical File: {params.step_1.autodesk_file_electrical.url.split('/')[-1]}"
            )

        doc.add_paragraph("")

        # Add legend
        doc.add_heading("Legend", level=2)
        legend_items = [
            ("✓ Green", "Category is in contract and present in model(s)"),
            ("✗ Orange", "Category is in contract but not in model(s)"),
            ("✗ Red", "Category is in model(s) but missing from contract"),
            ("✗ Gray", "Category is neither in contract nor in model(s)"),
        ]
        for symbol, description in legend_items:
            doc.add_paragraph(f"{symbol}: {description}", style="List Bullet")

        doc.add_paragraph("")

        # Add table
        doc.add_heading("Category Details", level=2)

        # Create table with 4 columns (added Element Count column)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Category"
        header_cells[1].text = "Status"
        header_cells[2].text = "Element Count"
        header_cells[3].text = "Description"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for category_name in all_categories:
            # Get element counts from both files
            structural_count = structural_counts.get(category_name, 0)
            electrical_count = electrical_counts.get(category_name, 0)
            total_count = structural_count + electrical_count

            # Check if category is in any model
            in_model = total_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Build element count display with breakdown
            if params.step_1.autodesk_file and params.step_1.autodesk_file_electrical:
                count_display = (
                    f"{total_count} (S:{structural_count}, E:{electrical_count})"
                )
            elif params.step_1.autodesk_file:
                count_display = f"{structural_count}"
            else:
                count_display = f"{electrical_count}"

            # Determine status symbol, description, and color
            if in_contract and in_model:
                status_symbol = "✓"
                status_text = "Present in contract and model(s)"
                color = RGBColor(0, 128, 0)  # Green
            elif in_contract and not in_model:
                status_symbol = "✗"
                status_text = "In contract but not in model(s)"
                color = RGBColor(255, 165, 0)  # Orange
            elif not in_contract and in_model:
                status_symbol = "✗"
                status_text = "Missing in the contract"
                color = RGBColor(255, 0, 0)  # Red
            else:  # not in_contract and not in_model
                status_symbol = "✗"
                status_text = "Not in contract, not in model(s)"
                color = RGBColor(128, 128, 128)  # Gray

            # Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = category_name
            row_cells[1].text = status_symbol
            row_cells[2].text = count_display
            row_cells[3].text = status_text

            # Apply color to status symbol
            for paragraph in row_cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Center align the element count
            for paragraph in row_cells[2].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vkt.progress_message("Finalizing document...", percentage=90)

        # Save document to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Create filename with timestamp
        filename = f"Contract_Compliance_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Return as DownloadResult
        return vkt.DownloadResult(vkt.File.from_data(doc_io.getvalue()), filename)
